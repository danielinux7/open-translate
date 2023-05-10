import os
import json
import requests
import sentencepiece as sp
from mosestokenizer import *
import re
from docx import Document

class Translator(object):
    
    
    def __init__(self, export_dir, language):
        self._tokenizer = sp.SentencePieceProcessor()
        self._detokenizer = sp.SentencePieceProcessor()
        self._tokenizer.load(os.path.join(export_dir, "assets", "src.model"))
        self._detokenizer.load(os.path.join(export_dir, "assets", "tgt.model"))
        self.lang = language
        
    def translate(self, texts):
        """Translates a batch of texts."""
        inputs, count = self._preprocess(texts)
        data = json.dumps({"inputs": inputs})
        headers = {"content-type": "application/json"}
        model_url = 'http://localhost:8501/v1/models/'+self.lang+':predict'
        json_response = requests.post(model_url, data=data, headers=headers)
        outputs = json.loads(json_response.text)['outputs']
        return self._postprocess(outputs,count)

    def _preprocess(self, texts):
        count = []
        tokenize = MosesTokenizer('ru')
        temp = []
        for text in texts:
           text = text.replace("-","- ")
           if text == '\r' or text == '':
             text = "▁"
           with MosesSentenceSplitter('ru') as splitsents:
              text = splitsents([text])
           count.append(len(text))
           temp.extend(text)
        texts = temp
        
        all_tokens = []
        lengths = []
        max_length = 0
        for text in texts:
            tokens = self._tokenizer.encode(text.lower(), out_type=str)
            tokens.insert(0,"<v7>")
            if self.lang == "ab-ru":
               tokens.insert(0,"<ab>")
            if self.lang == "ru-ab":
               tokens.insert(0,"<ru>")            
            length = len(tokens)
            all_tokens.append(tokens)
            lengths.append(length)
            max_length = max(max_length, length)
        for tokens, length in zip(all_tokens, lengths):
            if length < max_length:
                tokens += [""] * (max_length - length)

        inputs = {
            "tokens": all_tokens,
            "length": lengths
        }
        return inputs, count

    def _postprocess(self, outputs,count):
        texts = []
        for tokens, length in zip(outputs["tokens"], outputs["length"]):
            tokens = self._detokenizer.decode(tokens[0]).replace("- ","-")
            texts.append(tokens)
            
        temp = []
        i = 0    
        for length in count:
          text = ''
          for j in range(length):
            text = text.strip() + " " + texts[i+j].capitalize()
          temp.append(text.strip())
          i = i + length
        texts = temp
        return texts

def translate(src_list,sm_path,language):
    tgt_list = Translator(sm_path,language).translate(src_list)
    return tgt_list

def convFile(file):
    if file.filename[-5:] == ".docx":
        char_ab = {'ю':'ҩ', 'Ю':'Ҩ', '7':'ҵ', '?':'Ҵ', '6':'қ', ':':'Қ', '3':'ҷ', '№':'Ҷ', 'я':'ӷ', 'Я':'Ӷ','й':'ҟ', 'Й':'Ҟ', 
           'ё':'ӡ', 'Ё':'Ӡ', '8':'ԥ', '*':'Ԥ', '5':'џ', '%':'Џ', '0':'ҭ', ')':'Ҭ', '=':'ҿ', '+':'Ҿ', 'щ':'ҳ', 'Щ':'Ҳ', 
           'э':'ҽ', 'Э':'Ҽ', '9':')', 'ъ':'ә', 'Ъ':'Ә', '1':'?'}
        doc = Document(file)
        def replace(text):
            for key, value in char_ab.items():
              text = text.replace(key, value)
            return text

        def docx_replace_regex(doc_obj):

            for p in doc_obj.paragraphs:
                    inline = p.runs
                    # Loop added to work with runs (strings with same style)
                    for i in range(len(inline)):
                        if inline[i].font.name == "Arial Abkh":
                            text = replace(inline[i].text)
                            inline[i].text = text
                            inline[i].font.name = "Arial"

            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        docx_replace_regex(cell)
                        
        docx_replace_regex(doc)
        doc.save("./downloads/"+file.filename[:-5]+"_еиҭакны.docx")
        download = {'url':'/downloads/'+file.filename[:-5]+'_еиҭакны.docx','filename':file.filename[:-5]+'_еиҭакны.docx'}
        return download


